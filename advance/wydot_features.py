# wydot_features.py
# Minimalistic feature layer that plugs into your existing WYDOT Streamlit app.
# - Vector Fusion (Milvus + SQL rows)
# - Self-Validation Loop (LLM critique/revise)
# - Engineering Tools (MCP-style: math, unit conversion, quick formulas)
# - Workspace integration (uses your app’s helper if available)
# - RBAC + Audit trail over the same SQLite DB
# - Validator Dashboard (human vs AI)
# - Federated RAG (multi-collection simulation)
# - Knowledge Graph (synthetic) + simple graph reasoning
# - Agent Router (decide which tools to invoke)
# - Compliance Checker (toy rules)
# - Offline Cache (brute-force local embeddings)
# - Geo-RAG (folium HTML or fallback)
# - Report Generator (PDF/DOCX/Markdown fallback)

from __future__ import annotations
import os, io, json, time, math, random, sqlite3, re, pathlib
from typing import List, Dict, Any, Optional, Tuple, Callable
from dataclasses import dataclass, field
import difflib

# Optional deps (graceful fallback)
try:
    import sympy as sp
except Exception:
    sp = None

try:
    import folium
except Exception:
    folium = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except Exception:
    canvas = None

try:
    from docx import Document
except Exception:
    Document = None


# -----------------------------
# Utility
# -----------------------------
def _cosine(u: List[float], v: List[float]) -> float:
    if not u or not v or len(u) != len(v):
        return 0.0
    return sum(a*b for a,b in zip(u,v))

def _normed(vec: List[float]) -> List[float]:
    n = math.sqrt(sum(x*x for x in vec))
    return [x/n for x in vec] if n>0 else vec

def _strip(s: Optional[str]) -> str:
    return (s or "").strip()

def _now_ts() -> float:
    return time.time()

def _safe_json(val) -> str:
    try:
        return json.dumps(val, ensure_ascii=False)
    except Exception:
        return str(val)

def _ensure_dir(p: str):
    d = os.path.dirname(p)
    if d and not os.path.exists(d):
        os.makedirs(d, exist_ok=True)

# -----------------------------
# RBAC + Audit (uses same SQLite file)
# -----------------------------
class RBAC:
    def __init__(self, db_path: str):
        self.db_path = db_path
        self._conn = sqlite3.connect(db_path, check_same_thread=False, timeout=30)
        self._conn.execute("""
        CREATE TABLE IF NOT EXISTS roles (
            user_id INTEGER NOT NULL,
            role TEXT NOT NULL,
            granted_at REAL NOT NULL DEFAULT (strftime('%s','now')),
            PRIMARY KEY(user_id, role)
        )""")
        self._conn.execute("""
        CREATE TABLE IF NOT EXISTS audit_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            action TEXT NOT NULL,
            meta TEXT,
            ts REAL NOT NULL DEFAULT (strftime('%s','now'))
        )""")
        self._conn.commit()

    def grant(self, user_id: int, role: str):
        self._conn.execute("INSERT OR IGNORE INTO roles (user_id, role) VALUES (?,?)", (int(user_id), role))
        self._conn.commit()

    def revoke(self, user_id: int, role: str):
        self._conn.execute("DELETE FROM roles WHERE user_id=? AND role=?", (int(user_id), role))
        self._conn.commit()

    def has(self, user_id: int, role: str) -> bool:
        cur = self._conn.execute("SELECT 1 FROM roles WHERE user_id=? AND role=?", (int(user_id), role))
        return cur.fetchone() is not None

    def log(self, user_id: int, action: str, meta: Dict[str, Any] | None = None):
        self._conn.execute(
            "INSERT INTO audit_log (user_id, action, meta, ts) VALUES (?,?,?,?)",
            (int(user_id), action, _safe_json(meta or {}), _now_ts())
        )
        self._conn.commit()

    def list_audit(self, limit: int = 100) -> List[Dict[str, Any]]:
        cur = self._conn.execute("SELECT id,user_id,action,meta,ts FROM audit_log ORDER BY id DESC LIMIT ?", (limit,))
        rows = cur.fetchall()
        out = []
        for r in rows:
            out.append({"id": int(r[0]), "user_id": int(r[1]), "action": r[2], "meta": r[3], "ts": float(r[4])})
        return out


# -----------------------------
# Validator DB (human vs AI)
# -----------------------------
class ValidatorDB:
    def __init__(self, db_path: str):
        self._conn = sqlite3.connect(db_path, check_same_thread=False, timeout=30)
        self._conn.execute("""
        CREATE TABLE IF NOT EXISTS validation (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            question TEXT NOT NULL,
            ai_answer TEXT,
            human_answer TEXT,
            correct INTEGER NOT NULL DEFAULT 0,
            notes TEXT,
            ts REAL NOT NULL DEFAULT (strftime('%s','now'))
        )""")
        self._conn.commit()

    def add(self, user_id: int, question: str, ai_answer: str, human_answer: str, correct: bool, notes: str = ""):
        self._conn.execute(
            "INSERT INTO validation (user_id,question,ai_answer,human_answer,correct,notes,ts) VALUES (?,?,?,?,?,?,?)",
            (int(user_id), question, ai_answer, human_answer, 1 if correct else 0, notes, _now_ts())
        )
        self._conn.commit()

    def metrics(self) -> Dict[str, Any]:
        cur = self._conn.execute("SELECT COUNT(*), SUM(correct) FROM validation")
        n, s = cur.fetchone()
        n = int(n or 0)
        s = int(s or 0)
        acc = (s / n) if n > 0 else 0.0
        return {"total": n, "correct": s, "accuracy": acc}

    def recent(self, limit: int = 50) -> List[Dict[str, Any]]:
        cur = self._conn.execute("SELECT id,user_id,question,ai_answer,human_answer,correct,notes,ts FROM validation ORDER BY id DESC LIMIT ?", (limit,))
        out = []
        for r in cur.fetchall():
            out.append({
                "id": int(r[0]), "user_id": int(r[1]), "question": r[2],
                "ai_answer": r[3], "human_answer": r[4], "correct": bool(r[5]),
                "notes": r[6], "ts": float(r[7])
            })
        return out


# -----------------------------
# Vector Fusion: SQL-side tiny KB + embeddings
# -----------------------------
class VectorFusion:
    def __init__(self, db_path: str, embed_fn: Callable[[str], List[float]]):
        self._conn = sqlite3.connect(db_path, check_same_thread=False, timeout=30)
        self.embed_fn = embed_fn
        self._conn.execute("""
        CREATE TABLE IF NOT EXISTS sql_knowledge (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT,
            text TEXT NOT NULL,
            embedding TEXT
        )""")
        self._conn.commit()

    def seed_demo(self):
        cur = self._conn.execute("SELECT COUNT(*) FROM sql_knowledge")
        n = int(cur.fetchone()[0] or 0)
        if n > 0:
            return
        demo_rows = [
            ("Concrete Mix Limits", "For Class A concrete the slump shall be 50–100 mm and w/c <= 0.50."),
            ("Asphalt Compaction", "Target density shall be at least 92% of Gmm for WYDOT Section 401."),
            ("Steel Rebar Spacing", "Maximum spacing for #5 bars in bridge deck is 150 mm unless noted."),
            ("Work Zone Hours", "Lane closures on I-25 are restricted between 7am–9am and 4pm–6pm on weekdays."),
            ("Snow Fence Spec", "Timber snow fence posts shall conform to ASTM D1760 and be pressure treated."),
        ]
        for t, tx in demo_rows:
            self._conn.execute("INSERT INTO sql_knowledge (title, text) VALUES (?,?)", (t, tx))
        self._conn.commit()
        self.embed_all()

    def embed_all(self):
        cur = self._conn.execute("SELECT id, text FROM sql_knowledge WHERE embedding IS NULL OR embedding = ''")
        rows = cur.fetchall()
        for rid, text in rows:
            vec = self.embed_fn(text)
            self._conn.execute("UPDATE sql_knowledge SET embedding=? WHERE id=?", (json.dumps(vec), int(rid)))
        self._conn.commit()

    def search(self, query: str, k: int = 3) -> List[Dict[str, Any]]:
        qv = _normed(self.embed_fn(query))
        cur = self._conn.execute("SELECT id,title,text,embedding FROM sql_knowledge")
        scored = []
        for rid, title, txt, emb in cur.fetchall():
            try:
                v = json.loads(emb) if emb else []
            except Exception:
                v = []
            v = _normed(v)
            s = _cosine(qv, v)
            scored.append((s, {"title": title, "text": txt, "source": f"SQL:{title}"}))
        scored.sort(key=lambda x: x[0], reverse=True)
        return [x[1] for x in scored[:k]]

    def fuse(self, query: str, milvus_chunks: List[Dict[str, Any]], k_sql: int = 3, top_k: int = 6) -> List[Dict[str, Any]]:
        """
        Combine Milvus chunks (dicts with source/page/preview) and SQL rows.
        Ranking = simple alternate/score mix to show cross-domain context.
        """
        sql_hits = self.search(query, k=k_sql)
        fused = []
        i = j = 0
        while len(fused) < top_k and (i < len(milvus_chunks) or j < len(sql_hits)):
            if i < len(milvus_chunks):
                fused.append(milvus_chunks[i]); i += 1
            if len(fused) >= top_k: break
            if j < len(sql_hits):
                fused.append({
                    "doc_id": None,
                    "page": "-",
                    "source": sql_hits[j]["source"],
                    "preview": f"{sql_hits[j]['title']}: {sql_hits[j]['text'][:220]}",
                }); j += 1
        return fused


# -----------------------------
# Self-Validation (critique + revise)
# -----------------------------
def self_validate(client, model_id: str, question: str, draft: str, context: str = "") -> Tuple[str, str]:
    """
    Returns (revised_answer, critique_summary).
    """
    sys = (
        "You are a rigorous WYDOT QA reviewer. First list concrete issues with the draft "
        "(factuality, missing citations, clarity). Then produce a revised answer that is "
        "concise and grounded in the provided context."
    )
    prompt = (
        f"{sys}\n\n"
        f"Context:\n{context[:4000]}\n\n"
        f"Question: {question}\n\n"
        f"Draft answer:\n{draft}\n\n"
        "Output:\n1) Critique (3–6 bullet points)\n2) Revised answer (final)"
    )
    try:
        resp = client.models.generate_content(model=model_id, contents=[{"role":"user","parts":[{"text":prompt}]}])
        txt = (getattr(resp, "text", None) or "").strip()
        if not txt:
            return draft, "No critique produced."
        # crude split
        parts = re.split(r"\n\s*2\)\s*Revised answer", txt, flags=re.I)
        critique = parts[0].strip() if parts else "Critique unavailable."
        revised = txt if len(parts) < 2 else parts[1].strip()
        return (revised if revised else draft), critique
    except Exception as e:
        return draft, f"Critique failed: {e}"


# -----------------------------
# Engineering MCP-like tools
# -----------------------------
def solve_expression(expr: str) -> str:
    """
    SymPy-backed deterministic calculator. Examples:
      - '2*(3+4)^2'
      - 'solve(v-5, v)'
      - 'Reynolds(rho=1.225,V=20,L=0.5,mu=1.8e-5)'
    """
    if expr.strip().lower().startswith("reynolds"):
        m = dict(re.findall(r"([a-zA-Z]+)\s*=\s*([\-+eE0-9\.]+)", expr))
        try:
            rho = float(m.get("rho", "1.225"))
            V   = float(m.get("V", m.get("v","1.0")))
            L   = float(m.get("L", "1.0"))
            mu  = float(m.get("mu","1.8e-5"))
            Re  = rho*V*L/mu
            return f"Reynolds number ≈ {Re:.3e}"
        except Exception as e:
            return f"Error computing Reynolds: {e}"
    if sp is None:
        # super-simple safe eval for numbers only
        if not re.fullmatch(r"[0-9\.\+\-\*\/\(\)\s\^eE]+", expr):
            return "SymPy not installed and expression not purely numeric."
        expr = expr.replace("^", "**")
        try:
            return str(eval(expr, {"__builtins__":{}}))
        except Exception as e:
            return f"Eval error: {e}"
    try:
        expr2 = expr.replace("^", "**")
        # allow 'solve(x-2,x)' syntax
        if expr2.strip().lower().startswith("solve"):
            m = re.search(r"solve\((.+)\)", expr2, flags=re.I|re.S)
            if m:
                inside = m.group(1)
                parts = [p.strip() for p in inside.split(",")]
                if len(parts) == 2:
                    f_str, sym_str = parts
                    sym = sp.symbols(sym_str)
                    f = sp.sympify(f_str)
                    sol = sp.solve(sp.Eq(f,0), sym)
                    return f"Solutions for {sym_str}: {sol}"
        res = sp.sympify(expr2)
        return str(sp.N(res))
    except Exception as e:
        return f"SymPy error: {e}"

_UNIT_FACTORS = {
    # length (m)
    ("mm","m"): 1e-3, ("cm","m"): 1e-2, ("in","m"): 0.0254, ("ft","m"): 0.3048,
    ("m","mm"): 1e3, ("m","cm"): 1e2, ("m","in"): 39.37007874, ("m","ft"): 3.280839895,
    # speed
    ("km/h","m/s"): 1000/3600, ("m/s","km/h"): 3.6, ("mph","m/s"): 0.44704, ("m/s","mph"): 2.23693629,
    # force (N)
    ("kN","N"): 1000.0, ("N","kN"): 1/1000.0,
}

def convert_units(value: float, from_unit: str, to_unit: str) -> str:
    key = (from_unit.strip(), to_unit.strip())
    if key in _UNIT_FACTORS:
        return f"{value * _UNIT_FACTORS[key]:.6g} {to_unit}"
    return "Unsupported conversion."

# -----------------------------
# Federated RAG (multi-collection simulation)
# -----------------------------
# -----------------------------
# Enhanced Federated RAG (WYDOT Spec 2010 vs Latest)
# ================= Federated WYDOT (2010 vs 2021) =================


# Two live collections you mentioned
WYDOT_COLLECTIONS = ("wydotspec_llamaparse_2010", "wydotspec_llamaparse")

def _ensure_milvus_connected():
    """Connect to Zilliz/Milvus using env if not already connected."""
    try:
        from pymilvus import connections, utility
        try:
            # Will raise if no server
            _ = utility.get_server_version()
            return
        except Exception:
            uri = os.getenv("MILVUS_URI")
            token = os.getenv("MILVUS_TOKEN")
            if not (uri and token):
                raise RuntimeError("MILVUS_URI/MILVUS_TOKEN not set and no active connection.")
            connections.connect(alias="default", uri=uri, token=token)
    except Exception as e:
        raise RuntimeError(f"Milvus connection failed: {e}")

def _milvus_search_single(collection_name: str, query_vec: List[float], k: int = 5) -> List[Dict[str, Any]]:
    """Search one Milvus collection and return normalized hit dicts."""
    from pymilvus import Collection
    _ensure_milvus_connected()
    col = Collection(collection_name)
    col.load()
    # HNSW (COSINE) suggested params
    res = col.search(
        data=[query_vec],
        anns_field="vector",
        param={"metric_type": "COSINE", "params": {"ef": 64}},
        limit=k,
        output_fields=["doc_id", "chunk_id", "page", "section", "source", "content"],
    )
    hits = []
    version_label = "2010" if "2010" in collection_name else "2021"
    for hit in res[0]:
        fields = hit.fields
        hits.append({
            "collection": collection_name,
            "version": version_label,
            "score": float(hit.score),  # higher is better with COSINE
            "doc_id": fields.get("doc_id") or "",
            "page": int(fields.get("page") or -1),
            "section": fields.get("section") or "",
            "source": fields.get("source") or collection_name,
            "preview": (fields.get("content") or "")[:600]
        })
    return hits

def federated_search_wydot(query: str, embed_fn: Callable[[str], List[float]], k_each: int = 5) -> Dict[str, List[Dict[str, Any]]]:
    """
    Retrieve top-k from both versions independently.
    Returns dict: {"2010": [...], "2021": [...], "all": [...]}
    """
    qv = _normed(embed_fn(query))
    hits_2010 = _milvus_search_single(WYDOT_COLLECTIONS[0], qv, k=k_each)
    hits_2021 = _milvus_search_single(WYDOT_COLLECTIONS[1], qv, k=k_each)
    # Interleave for a quick fused view
    fused = []
    for i in range(max(len(hits_2010), len(hits_2021))):
        if i < len(hits_2010): fused.append(hits_2010[i])
        if i < len(hits_2021): fused.append(hits_2021[i])
    return {"2010": hits_2010, "2021": hits_2021, "all": fused}

def _pair_hits_by_similarity(hits_a: List[Dict[str, Any]], hits_b: List[Dict[str, Any]], min_sim: float = 0.30) -> List[Dict[str, Any]]:
    """
    Greedy pairings between 2010 and 2021 chunks based on text similarity (SequenceMatcher ratio).
    Returns list of {"sim": float, "a": hit2010|None, "b": hit2021|None}
    """
    used_b = set()
    pairs = []
    for a in hits_a:
        best_j, best_sim = None, 0.0
        for j, b in enumerate(hits_b):
            if j in used_b: 
                continue
            r = difflib.SequenceMatcher(None, a["preview"], b["preview"]).ratio()
            if r > best_sim:
                best_sim, best_j = r, j
        if best_j is not None and best_sim >= min_sim:
            pairs.append({"sim": best_sim, "a": a, "b": hits_b[best_j]})
            used_b.add(best_j)
        else:
            pairs.append({"sim": 0.0, "a": a, "b": None})
    # Any leftover 2021 hits with no pair
    for j, b in enumerate(hits_b):
        if j not in used_b:
            pairs.append({"sim": 0.0, "a": None, "b": b})
    # strongest pairs first
    pairs.sort(key=lambda x: x["sim"], reverse=True)
    return pairs

def _summarize_pairs_programmatically(pairs: List[Dict[str, Any]], max_pairs: int = 5) -> str:
    """
    Heuristic, no-LLM summary highlighting overlaps and differences.
    """
    if not pairs:
        return "No comparable passages found between 2010 and 2021 for this query."

    lines = []
    covered = 0
    for p in pairs:
        if covered >= max_pairs:
            break
        a, b, sim = p.get("a"), p.get("b"), p.get("sim", 0.0)
        if a and b:
            lines.append(
                f"- Overlap (sim≈{sim:.2f}): 2010 p.{a['page']} vs 2021 p.{b['page']} — "
                f"both mention: { _common_phrase(a['preview'], b['preview']) }"
            )
        elif a and not b:
            lines.append(f"- 2010-only: p.{a['page']} — {a['preview'][:140]}…")
        elif b and not a:
            lines.append(f"- 2021-only: p.{b['page']} — {b['preview'][:140]}…")
        covered += 1
    return "Comparative summary:\n" + "\n".join(lines)

def _common_phrase(t1: str, t2: str, min_tokens: int = 3) -> str:
    """
    Tiny helper to surface a plausible shared phrase by longest common substring on tokens.
    """
    a = t1.split()
    b = t2.split()
    # map words to indices
    idx = {}
    best_len, best_end = 0, 0
    prev = {}
    for i, w1 in enumerate(a, 1):
        curr = {}
        for j, w2 in enumerate(b, 1):
            if w1 == w2:
                curr[j] = prev.get(j-1, 0) + 1
                if curr[j] > best_len:
                    best_len, best_end = curr[j], i
        prev = curr
    if best_len >= min_tokens:
        start = best_end - best_len
        return " ".join(a[start:best_end])[:120]
    # fallback: first 6 tokens of t1
    return " ".join(a[:6])[:60]

# def federated_search(
#     query: str,
#     embed_fn: Callable[[str], List[float]],
#     k_each: int = 5,
#     sim_min: float = 0.30,
#     llm_client=None,
#     llm_model: Optional[str] = None,
# ) -> Dict[str, Any]:
#     """
#     One-shot federated retrieval + comparison for WYDOT specs.
#     Returns:
#       {
#         "pairs": [...],
#         "summary": "...programmatic or LLM diff...",
#         "top_context": "concatenated snippets for optional LLM prompting",
#         "answer_2010": "...answer using only 2010 spec passages...",
#         "answer_latest": "...answer using only 2021 spec passages...",
#         "gemini_diff_summary": "...same as summary (LLM diff)...",
#         "federated": [list of combined hits from both versions]
#       }
#     """
#     # 1) Retrieve from both editions
#     res = federated_search_wydot(query, embed_fn, k_each=k_each)
#     hits_2010 = res["2010"]
#     hits_2021 = res["2021"]

#     # 2) Pair by similarity
#     pairs = _pair_hits_by_similarity(hits_2010, hits_2021, min_sim=sim_min)

#     # 3) Build contexts:
#     #    - top_context: mixed 2010/2021 for diffing
#     #    - ctx2010: 2010-only
#     #    - ctx_latest: 2021-only
#     ctx_lines: List[str] = []
#     ctx2010_lines: List[str] = []
#     ctx_latest_lines: List[str] = []

#     max_pairs_for_ctx = max(6, k_each)
#     for p in pairs[:max_pairs_for_ctx]:
#         a, b = p.get("a"), p.get("b")
#         if a:
#             line_2010 = f"[2010 p.{a['page']}] {a['preview']}"
#             ctx_lines.append(line_2010)
#             ctx2010_lines.append(line_2010)
#         if b:
#             line_2021 = f"[2021 p.{b['page']}] {b['preview']}"
#             ctx_lines.append(line_2021)
#             ctx_latest_lines.append(line_2021)

#     top_context = "\n".join(ctx_lines)[:6000]
#     ctx2010 = "\n".join(ctx2010_lines)[:6000]
#     ctx_latest = "\n".join(ctx_latest_lines)[:6000]

#     # 4) Programmatic summary first
#     summary = _summarize_pairs_programmatically(pairs)
#     gemini_diff_summary = summary

#     # 5) If an LLM client is provided, ask it to produce a concise diff with citations
#     if llm_client and llm_model and top_context:
#         try:
#             prompt = (
#                 "You are comparing two editions of the WYDOT Standard Specifications (2010 vs 2021).\n"
#                 "Given the paired excerpts below, produce a concise answer to the user query, explicitly noting:\n"
#                 " - where both editions agree,\n"
#                 " - where the 2021 edition tightens/relaxes/clarifies requirements,\n"
#                 " - and include inline citations like [2010 p.X] or [2021 p.Y] after each claim.\n\n"
#                 f"Query: {query}\n\n"
#                 f"Paired context:\n{top_context}\n\n"
#                 "Answer:"
#             )
#             resp = llm_client.models.generate_content(
#                 model=llm_model,
#                 contents=[{"role": "user", "parts": [{"text": prompt}]}],
#             )
#             llm_txt = (getattr(resp, "text", "") or "").strip()
#             if llm_txt:
#                 summary = llm_txt
#                 gemini_diff_summary = llm_txt
#         except Exception as e:
#             summary = summary + f"\n\n(LLM summarization unavailable: {e})"
#             gemini_diff_summary = summary

#     # 6) Per-spec answers (2010-only and 2021-only)
#     answer_2010 = ""
#     answer_latest = ""

#     if llm_client and llm_model:
#         # 2010-only
#         if ctx2010:
#             try:
#                 prompt_2010 = (
#                     "You are a WYDOT specifications assistant.\n"
#                     "Using ONLY the 2010 specification excerpts below, answer the user's question.\n"
#                     "If the answer is not clearly stated in these passages, say that the 2010 excerpts shown do not fully answer it.\n\n"
#                     f"User question:\n{query}\n\n"
#                     f"2010 context:\n{ctx2010}\n\n"
#                     "Answer (2010 spec only):"
#                 )
#                 resp_2010 = llm_client.models.generate_content(
#                     model=llm_model,
#                     contents=[{"role": "user", "parts": [{"text": prompt_2010}]}],
#                 )
#                 answer_2010 = (getattr(resp_2010, "text", "") or "").strip()
#             except Exception as e:
#                 answer_2010 = f"(2010 answer unavailable: {e})"

#         # 2021-only (latest)
#         if ctx_latest:
#             try:
#                 prompt_latest = (
#                     "You are a WYDOT specifications assistant.\n"
#                     "Using ONLY the 2021 (latest) specification excerpts below, answer the user's question.\n"
#                     "If the answer is not clearly stated in these passages, say that the 2021 excerpts shown do not fully answer it.\n\n"
#                     f"User question:\n{query}\n\n"
#                     f"2021 context:\n{ctx_latest}\n\n"
#                     "Answer (latest spec only):"
#                 )
#                 resp_latest = llm_client.models.generate_content(
#                     model=llm_model,
#                     contents=[{"role": "user", "parts": [{"text": prompt_latest}]}],
#                 )
#                 answer_latest = (getattr(resp_latest, "text", "") or "").strip()
#             except Exception as e:
#                 answer_latest = f"(Latest spec answer unavailable: {e})"

#     # 7) Flatten combined hits for convenience in the UI
#     federated: List[Dict[str, Any]] = []
#     for p in pairs:
#         a, b = p.get("a"), p.get("b")
#         if a:
#             federated.append(a)
#         if b:
#             federated.append(b)
#     # print(answer_2010, answer_latest)
#     return {
#         "pairs": pairs,
#         "summary": summary,
#         "top_context": top_context,
#         "answer_2010": answer_2010,
#         "answer_latest": answer_latest,
#         "gemini_diff_summary": gemini_diff_summary,
#         "federated": federated,
#     }
def federated_search(
    query: str,
    embed_fn: Callable[[str], List[float]],
    k_each: int = 5,
    sim_min: float = 0.30,  # kept for signature compatibility, not used
    llm_client=None,
    llm_model: Optional[str] = None,
    history_text_2010: str = "(no previous turns)",   # 👈 NEW ARG
) -> Dict[str, Any]:
    """
    2010-only WYDOT spec answer.

    - Uses ONLY the 2010 vector store (WYDOT_COLLECTIONS[0]).
    - Returns a single grounded answer: {"answer_2010": "..."}.
    - Uses its own 2010-only chat history passed in as `history_text_2010`.
    """
    if llm_client is None or llm_model is None:
        raise ValueError("llm_client and llm_model are required to generate answer_2010")

    # --- 1) Retrieve from 2010 collection only ---
    qv = _normed(embed_fn(query))
    hits_2010 = _milvus_search_single(WYDOT_COLLECTIONS[0], qv, k=k_each)

    if not hits_2010:
        return {
            "answer_2010": (
                "I could not find any 2010 WYDOT specification excerpts in the vector store "
                "that clearly answer this question."
            )
        }

    # --- 2) Build 2010 context ---
    ctx_lines = []
    for h in hits_2010:
        line = f"[2010 p.{h['page']}] {h['preview']}"
        ctx_lines.append(line)

    ctx2010 = "\n".join(ctx_lines)[:6000]

    # --- 3) Prompt: use 2010-only history ---
    base = (
        "You are WYDOT chatbot, a polite and helpful Virtual Assistant of Wyoming Department of Transportation (WYDOT).\n"
        "Answer the question using the provided context and (if relevant) the prior conversation.\n\n"
        "Conversation so far (most recent turns first):\n"
        "{history}\n\n"
        "Context inside double backticks:``{context}``\n"
    )

    tail = (
        "Question inside triple backticks:```{question}```\n"
        "If the question is out of scope, answer based on your role.\n"
        "JUST PROVIDE THE ANSWER IN ENGLISH WITHOUT ``` AND NOTHING ELSE.\n"
    )

    prompt_2010 = base.format(
        history=history_text_2010,
        context=ctx2010,
    ) + tail.format(question=query)

    # --- 4) Call Gemini ---
    try:
        resp_2010 = llm_client.models.generate_content(
            model=llm_model,
            contents=[{"role": "user", "parts": [{"text": prompt_2010}]}],
        )
        answer_2010 = (getattr(resp_2010, "text", "") or "").strip()
    except Exception as e:
        answer_2010 = f"(2010 answer unavailable due to error: {e})"

    return {
        "answer_2010": answer_2010
    }





# -----------------------------
# Knowledge Graph (synthetic)
# -----------------------------
@dataclass
class KG:
    edges: Dict[str, List[str]] = field(default_factory=dict)

    def add_edge(self, a: str, b: str):
        self.edges.setdefault(a, []).append(b)
        self.edges.setdefault(b, []).append(a)

    def build_demo(self):
        pairs = [
            ("Section 401 Asphalt", "Gmm Density"),
            ("Section 401 Asphalt", "Compaction Target"),
            ("Section 501 Concrete", "Slump Range"),
            ("Section 501 Concrete", "w/c Limit"),
            ("ASTM D1760", "Snow Fence"),
            ("Bridge Deck", "Rebar Spacing"),
        ]
        for a,b in pairs:
            self.add_edge(a,b)

    def related(self, term: str) -> List[str]:
        term = term.strip()
        out = []
        for k,v in self.edges.items():
            if term.lower() in k.lower():
                out.extend(v)
        return sorted(set(out))

# -----------------------------
# Compliance Checker (toy)
# -----------------------------
_SPEC_RULES = {
    "slump_mm_min": 50.0,
    "slump_mm_max": 100.0,
    "wc_max": 0.50,
    "rebar_spacing_mm_max": 150.0,
    "asphalt_gmm_density_min": 0.92,  # 92%
}

def check_compliance(text: str) -> List[str]:
    """
    Parse numbers from user text and check against simple rules.
    Accepts snippets like: "slump 85 mm, w/c 0.48, spacing 160 mm"
    """
    notes = []
    t = text.lower()
    slump = re.search(r"slump\s*([\d\.]+)\s*mm", t)
    wc    = re.search(r"w/?c\s*([\d\.]+)", t)
    spc   = re.search(r"(spacing|rebar)\s*([\d\.]+)\s*mm", t)
    gmm   = re.search(r"(gmm|density)\s*([\d\.]+)", t)

    if slump:
        val = float(slump.group(1))
        if val < _SPEC_RULES["slump_mm_min"] or val > _SPEC_RULES["slump_mm_max"]:
            notes.append(f"❌ Slump {val} mm is out of 50–100 mm")
        else:
            notes.append(f"✅ Slump {val} mm OK")
    if wc:
        val = float(wc.group(1))
        if val > _SPEC_RULES["wc_max"]:
            notes.append(f"❌ w/c {val} exceeds {_SPEC_RULES['wc_max']}")
        else:
            notes.append(f"✅ w/c {val} OK")
    if spc:
        val = float(spc.group(2))
        if val > _SPEC_RULES["rebar_spacing_mm_max"]:
            notes.append(f"❌ Rebar spacing {val} mm > 150 mm")
        else:
            notes.append(f"✅ Rebar spacing {val} mm OK")
    if gmm:
        val = float(gmm.group(2))
        if val < _SPEC_RULES["asphalt_gmm_density_min"]:
            notes.append(f"❌ Asphalt density {val:.2f} < 0.92 Gmm")
        else:
            notes.append(f"✅ Asphalt density {val:.2f} OK")
    if not notes:
        notes.append("No recognizable design inputs found.")
    return notes

# -----------------------------
# Offline Cache (JSON embeddings)
# -----------------------------
class OfflineCache:
    """
    Offline RAG using CURRENT session conversation as the document corpus.
    Every user + assistant message becomes a retrievable chunk.
    """

    def __init__(self, embed_fn: Callable[[str], List[float]]):
        self.embed_fn = embed_fn
        self.items: List[Dict[str, Any]] = []

    def load_from_session(self, messages: List[Dict[str, Any]]):
        """
        Convert chat history into RAG chunks.
        messages = [
            {"role": "user", "content": "..."},
            {"role": "assistant", "content": "..."}
        ]
        """
        self.items = []
        for idx, msg in enumerate(messages):
            text = msg.get("content") or ""
            if not text.strip():
                continue

            emb = self.embed_fn(text)

            self.items.append({
                "doc_id": "session_doc",
                "chunk_id": idx,
                "role": msg.get("role"),
                "page": idx + 1,         # Virtual page number
                "section": msg.get("role"),
                "source": "offline:chat_history",
                "content": text,
                "preview": text[:300],
                "emb": emb,
            })

    def search(self, query: str, k: int = 5) -> List[Dict[str, Any]]:
        qv = _normed(self.embed_fn(query))
        scored = []

        for d in self.items:
            sim = _cosine(qv, _normed(d["emb"]))
            d2 = dict(d)
            d2["score"] = sim       # makes “confidence % chips” work
            scored.append((sim, d2))

        scored.sort(key=lambda x: x[0], reverse=True)
        return [d for _, d in scored[:k]]


# -----------------------------
# Geo-RAG (folium or fallback)
# -----------------------------
def render_map_html(points: List[Tuple[float,float,str]]) -> str:
    if folium:
        lat0, lon0 = points[0][0], points[0][1]
        m = folium.Map(location=[lat0,lon0], zoom_start=6)
        for lat,lon,label in points:
            folium.Marker([lat,lon], popup=label).add_to(m)
        return m._repr_html_()
    # Fallback tiny HTML map (very basic)
    lis = "".join([f"<li>{round(lat,4)}, {round(lon,4)} — {label}</li>" for lat,lon,label in points])
    return f"<div><h4>Map (fallback)</h4><ul>{lis}</ul></div>"

# -----------------------------
# Report generator
# -----------------------------
def make_report(question: str, answer: str, sources: List[Dict[str, Any]], metrics: Dict[str, Any], out_dir: str = "./reports") -> Dict[str, str]:
    os.makedirs(out_dir, exist_ok=True)
    ts = int(_now_ts())
    base = f"wydot_report_{ts}"
    md_path = os.path.join(out_dir, base + ".md")

    lines = [f"# WYDOT Answer Report\n\n**Question:** {question}\n\n**Answer:**\n\n{answer}\n\n## Sources\n"]
    for i,s in enumerate(sources,1):
        lines.append(f"{i}. {s.get('source')} (page {s.get('page')}) — {s.get('preview','')[:140]}")

    if metrics:
        lines.append("\n## Metrics\n")
        for k,v in metrics.items():
            lines.append(f"- {k}: {v}")

    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    outputs = {"markdown": md_path}

    # Optional PDF
    if canvas:
        pdf_path = os.path.join(out_dir, base + ".pdf")
        c = canvas.Canvas(pdf_path, pagesize=letter)
        width, height = letter
        y = height - 50
        def draw_line(txt: str):
            nonlocal y
            for seg in re.findall(".{1,95}", txt):
                c.drawString(40, y, seg)
                y -= 14
                if y < 60:
                    c.showPage(); y = height - 50
        draw_line(f"WYDOT Report — {time.ctime(ts)}")
        draw_line(f"Question: {question}")
        draw_line("Answer:")
        draw_line(answer)
        draw_line("Sources:")
        for s in sources:
            draw_line(f"- {s.get('source')} page {s.get('page')} — {s.get('preview','')[:100]}")
        if metrics:
            draw_line("Metrics:")
            for k,v in metrics.items():
                draw_line(f"- {k}: {v}")
        c.save()
        outputs["pdf"] = pdf_path

    # Optional DOCX
    if Document:
        docx_path = os.path.join(out_dir, base + ".docx")
        doc = Document()
        doc.add_heading("WYDOT Answer Report", level=1)
        doc.add_paragraph(f"Question: {question}")
        doc.add_heading("Answer", level=2)
        doc.add_paragraph(answer)
        doc.add_heading("Sources", level=2)
        for s in sources:
            doc.add_paragraph(f"- {s.get('source')} page {s.get('page')} — {s.get('preview','')[:100]}")
        if metrics:
            doc.add_heading("Metrics", level=2)
            for k,v in metrics.items():
                doc.add_paragraph(f"- {k}: {v}")
        doc.save(docx_path)
        outputs["docx"] = docx_path

    return outputs

# -----------------------------
# Agent Router (toy)
# -----------------------------
def route(query: str) -> str:
    q = query.lower()
    if any(w in q for w in ["convert","unit","km/h","mph","m/s"]):
        return "units"
    if any(w in q for w in ["solve","reynolds","equation","compute"]):
        return "calc"
    if any(w in q for w in ["map","location","lat","lon","route"]):
        return "geo"
    if any(w in q for w in ["compliance","within spec","limit"]):
        return "compliance"
    return "rag"

# -----------------------------
# Workspace helper passthrough
# -----------------------------
def workspace_summarize(get_workspace_context_fn: Callable[[str,int], str], query: str) -> str:
    try:
        ctx = get_workspace_context_fn(query, max_chars=2000)
        return ctx or "(No Workspace snippets found.)"
    except Exception as e:
        return f"(Workspace unavailable: {e})"
